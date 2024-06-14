from pdf2image import convert_from_path
from docx import Document
from PIL import Image
import io


def pdf_to_images(pdf_file):
    images = convert_from_path(pdf_file)
    return images


def save_images_to_docx(images, docx_file):
    doc = Document()
    for image in images:
        img_io = io.BytesIO()
        image.save(img_io, format='PNG')
        doc.add_picture(img_io, width=doc.shared.Inches(5))
        doc.add_paragraph()
    doc.save(docx_file)


if __name__ == "__main__":
    pdf_file = r"C:\Users\fengk\Desktop\TPA产品服务说明.pdf"
    docx_file = r"C:\Users\fengk\Desktop\TPA产品服务说明.docx"

    images = pdf_to_images(pdf_file)
    save_images_to_docx(images, docx_file)

    print("转换完成！")
